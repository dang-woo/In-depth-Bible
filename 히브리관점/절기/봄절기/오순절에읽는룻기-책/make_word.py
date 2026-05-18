import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"C:\obsidian\dang-woo\히브리관점\절기\봄절기\오순절에읽는룻기-책"

FILES = [
    "00-목차.md",
    "01-개요-왜-오순절에-룻기를-읽는가.md",
    "02-룻기-1장-모압에서-베들레헴으로.md",
    "03-룻기-2장-우연처럼-찾아온-고엘.md",
    "04-룻기-3장-타작마당의-밤.md",
    "05-룻기-4장-성문에서-완성된-구속.md",
    "06-오순절과-예수님-말씀과-성령의-연합.md",
    "07-핵심-히브리어-단어-정리.md",
]

def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    rPr.insert(0, rFonts)

def set_para_font(para, size=11):
    for run in para.runs:
        set_font(run, size)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    return br

def add_page_break_v2(doc):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

def strip_obsidian(text):
    # Remove [[link|display]] -> display
    text = re.sub(r'\[\[.*?\|(.*?)\]\]', r'\1', text)
    # Remove [[link]] -> link
    text = re.sub(r'\[\[(.*?)\]\]', r'\1', text)
    return text

def is_nav_line(line):
    stripped = line.strip()
    if not stripped:
        return False
    # Lines that are pure navigation (contain links and navigation words)
    if re.match(r'^(이전 장|다음 장|목차로)', stripped):
        return True
    return False

def add_runs_with_bold(para, text, base_size=11, italic=False):
    """Parse **bold** and add runs accordingly."""
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        is_bold = (i % 2 == 1)
        set_font(run, size=base_size, bold=is_bold, italic=italic)

def process_file(doc, filepath, is_first=False):
    with open(filepath, encoding='utf-8') as f:
        raw = f.read()

    lines = raw.split('\n')
    i = 0
    in_blockquote = False
    bq_lines = []

    def flush_blockquote():
        nonlocal in_blockquote, bq_lines
        if bq_lines:
            combined = ' '.join(bq_lines).strip()
            # Remove leading > markers
            combined = re.sub(r'^>\s*', '', combined)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.2)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            # Add left border via XML
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '4')
            left.set(qn('w:space'), '4')
            left.set(qn('w:color'), '888888')
            pBdr.append(left)
            pPr.append(pBdr)
            add_runs_with_bold(p, combined, base_size=10, italic=True)
            for run in p.runs:
                set_font(run, size=10, italic=True, bold=run.bold, color=(80,80,80))
        bq_lines = []
        in_blockquote = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip navigation lines
        if is_nav_line(stripped):
            i += 1
            continue

        # H1
        if line.startswith('# ') and not line.startswith('## '):
            flush_blockquote()
            text = strip_obsidian(line[2:].strip())
            p = doc.add_heading(text, level=1)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(12)
            for run in p.runs:
                set_font(run, size=18, bold=True)
            i += 1
            continue

        # H2
        if line.startswith('## '):
            flush_blockquote()
            text = strip_obsidian(line[3:].strip())
            p = doc.add_heading(text, level=2)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                set_font(run, size=13, bold=True)
            i += 1
            continue

        # H3
        if line.startswith('### '):
            flush_blockquote()
            text = strip_obsidian(line[4:].strip())
            p = doc.add_heading(text, level=3)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                set_font(run, size=11, bold=True)
            i += 1
            continue

        # Horizontal rule
        if stripped in ('---', '***', '___'):
            flush_blockquote()
            p = doc.add_paragraph('─' * 40)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                set_font(run, size=9, color=(180,180,180))
            i += 1
            continue

        # Table row (simple handling)
        if stripped.startswith('|'):
            flush_blockquote()
            # Skip separator rows
            if re.match(r'^\|[-| :]+\|$', stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.strip('|').split('|')]
            # Detect if header (next line is separator)
            next_line = lines[i+1].strip() if i+1 < len(lines) else ''
            is_header = bool(re.match(r'^\|[-| :]+\|$', next_line))
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            text = '  |  '.join(cells)
            add_runs_with_bold(p, text, base_size=10)
            if is_header:
                for run in p.runs:
                    run.bold = True
                    set_font(run, size=10, bold=True)
            else:
                for run in p.runs:
                    set_font(run, size=10)
            i += 1
            continue

        # Blockquote
        if line.startswith('>'):
            content = line[1:].strip()
            # Continuation of blockquote
            if content or in_blockquote:
                in_blockquote = True
                if content:
                    bq_lines.append(content)
                else:
                    flush_blockquote()
            i += 1
            continue
        else:
            if in_blockquote:
                flush_blockquote()

        # Empty line
        if not stripped:
            i += 1
            continue

        # Regular paragraph
        text = strip_obsidian(stripped)
        if text:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Cm(0)
            add_runs_with_bold(p, text, base_size=11)

        i += 1

    flush_blockquote()


def setup_styles(doc):
    # Page margins
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Normal style
    normal = doc.styles['Normal']
    normal.font.name = '맑은 고딕'
    normal.font.size = Pt(11)
    from docx.oxml.ns import qn
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = '맑은 고딕'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    h1.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = '맑은 고딕'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    h2.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = '맑은 고딕'
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    h3.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')


def main():
    doc = Document()
    setup_styles(doc)

    # Cover-style title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    run = p.add_run("오순절에 읽는 룻기")
    set_font(run, size=24, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(16)
    run2 = p2.add_run("보리 추수에서 밀 추수까지,\n룻의 선택에서 예수님의 고엘 사역까지 이어지는 이야기")
    set_font(run2, size=12, color=(80, 80, 80), italic=True)

    for f in FILES:
        add_page_break_v2(doc)
        fpath = os.path.join(BASE, f)
        if os.path.exists(fpath):
            process_file(doc, fpath)
        else:
            print(f"  [건너뜀] {f} 없음")

    out = os.path.join(BASE, "오순절에읽는룻기.docx")
    doc.save(out)
    print(f"\n저장 완료: {out}")

if __name__ == "__main__":
    main()
